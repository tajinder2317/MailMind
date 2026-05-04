"""
Attachment Processing Worker for MailMind

This module handles text extraction from various file formats including
PDF, DOCX, and images (PNG/JPG) using OCR capabilities.
"""

import asyncio
import io
import logging
from typing import Optional, Dict, Any
from pathlib import Path
import tempfile
import os

# PDF processing
import PyPDF2
import pdfplumber

# DOCX processing
from docx import Document

# Image OCR processing
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR not available. Install pytesseract and Pillow for image text extraction.")

logger = logging.getLogger(__name__)


class AttachmentProcessor:
    """
    Processes email attachments and extracts text content.
    
    Supported formats:
    - PDF files (using PyPDF2 and pdfplumber)
    - DOCX files (using python-docx)
    - Image files (PNG/JPG using Tesseract OCR)
    """
    
    SUPPORTED_MIME_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'image/png': 'image',
        'image/jpeg': 'image',
        'image/jpg': 'image',
    }
    
    def __init__(self, ocr_language: str = 'eng'):
        """
        Initialize attachment processor.
        
        Args:
            ocr_language: Language code for OCR (default: 'eng' for English)
        """
        self.ocr_language = ocr_language
        
        # Configure Tesseract path if needed (Windows)
        if OCR_AVAILABLE and os.name == 'nt':  # Windows
            tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
            if os.path.exists(tesseract_path):
                pytesseract.pytesseract.tesseract_cmd = tesseract_path
            else:
                logger.warning("Tesseract not found at default Windows path")
    
    async def extract_text(
        self, 
        file_content: bytes, 
        filename: str, 
        mime_type: str
    ) -> Dict[str, Any]:
        """
        Extract text content from attachment.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            mime_type: MIME type of the file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            file_type = self.SUPPORTED_MIME_TYPES.get(mime_type.lower())
            
            if not file_type:
                return {
                    'status': 'failed',
                    'error': f'Unsupported file type: {mime_type}',
                    'text': None
                }
            
            logger.info(f"Processing {file_type} file: {filename}")
            
            # Route to appropriate processor
            if file_type == 'pdf':
                result = await self._extract_pdf_text(file_content, filename)
            elif file_type == 'docx':
                result = await self._extract_docx_text(file_content, filename)
            elif file_type == 'image':
                result = await self._extract_image_text(file_content, filename)
            else:
                result = {
                    'status': 'failed',
                    'error': f'Unsupported file type: {file_type}',
                    'text': None
                }
            
            logger.info(f"Successfully processed {filename}: {result['status']}")
            return result
            
        except Exception as e:
            logger.error(f"Error processing attachment {filename}: {str(e)}")
            return {
                'status': 'failed',
                'error': str(e),
                'text': None
            }
    
    async def _extract_pdf_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from PDF file using multiple methods.
        
        Args:
            file_content: PDF file bytes
            filename: Original filename
            
        Returns:
            Dictionary with extracted text and metadata
        """
        text_content = ""
        extraction_method = None
        page_count = 0
        
        try:
            # Method 1: Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                extraction_method = "pdfplumber"
            
            # Fallback to PyPDF2 if pdfplumber didn't extract much text
            if len(text_content.strip()) < 100:
                text_content = ""
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                page_count = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                
                extraction_method = "PyPDF2"
            
            if not text_content.strip():
                return {
                    'status': 'failed',
                    'error': 'No text could be extracted from PDF',
                    'text': None,
                    'metadata': {
                        'page_count': page_count,
                        'extraction_method': None
                    }
                }
            
            return {
                'status': 'completed',
                'text': text_content.strip(),
                'metadata': {
                    'page_count': page_count,
                    'extraction_method': extraction_method,
                    'character_count': len(text_content)
                }
            }
            
        except Exception as e:
            logger.error(f"PDF extraction failed for {filename}: {str(e)}")
            return {
                'status': 'failed',
                'error': f'PDF extraction error: {str(e)}',
                'text': None
            }
    
    async def _extract_docx_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: DOCX file bytes
            filename: Original filename
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            # Create temporary file for docx processing
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Process with python-docx
                doc = Document(temp_file_path)
                
                # Extract text from paragraphs
                paragraphs = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text)
                
                # Extract text from tables
                table_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            table_text.append(' | '.join(row_text))
                
                # Combine all text
                all_text = '\n'.join(paragraphs)
                if table_text:
                    all_text += '\n\n--- Tables ---\n' + '\n'.join(table_text)
                
                if not all_text.strip():
                    return {
                        'status': 'failed',
                        'error': 'No text could be extracted from DOCX',
                        'text': None
                    }
                
                return {
                    'status': 'completed',
                    'text': all_text.strip(),
                    'metadata': {
                        'paragraph_count': len(paragraphs),
                        'table_count': len(doc.tables),
                        'character_count': len(all_text)
                    }
                }
                
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
                
        except Exception as e:
            logger.error(f"DOCX extraction failed for {filename}: {str(e)}")
            return {
                'status': 'failed',
                'error': f'DOCX extraction error: {str(e)}',
                'text': None
            }
    
    async def _extract_image_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from image using OCR.
        
        Args:
            file_content: Image file bytes
            filename: Original filename
            
        Returns:
            Dictionary with extracted text and metadata
        """
        if not OCR_AVAILABLE:
            return {
                'status': 'failed',
                'error': 'OCR not available. Install pytesseract and Pillow.',
                'text': None
            }
        
        try:
            # Open image with PIL
            image = Image.open(io.BytesIO(file_content))
            
            # Preprocess image for better OCR
            image = self._preprocess_image(image)
            
            # Extract text using Tesseract
            text_content = pytesseract.image_to_string(
                image, 
                lang=self.ocr_language,
                config='--psm 6'  # Assume uniform block of text
            )
            
            # Also try with different page segmentation modes if needed
            if len(text_content.strip()) < 10:
                # Try with PSM 3 (fully automatic)
                alt_text = pytesseract.image_to_string(
                    image,
                    lang=self.ocr_language,
                    config='--psm 3'
                )
                if len(alt_text.strip()) > len(text_content.strip()):
                    text_content = alt_text
            
            if not text_content.strip():
                return {
                    'status': 'completed',
                    'text': '',
                    'metadata': {
                        'ocr_method': 'tesseract',
                        'character_count': 0,
                        'note': 'No text detected in image'
                    }
                }
            
            return {
                'status': 'completed',
                'text': text_content.strip(),
                'metadata': {
                    'ocr_method': 'tesseract',
                    'character_count': len(text_content),
                    'image_size': image.size,
                    'image_mode': image.mode
                }
            }
            
        except Exception as e:
            logger.error(f"Image OCR failed for {filename}: {str(e)}")
            return {
                'status': 'failed',
                'error': f'OCR extraction error: {str(e)}',
                'text': None
            }
    
    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """
        Preprocess image for better OCR results.
        
        Args:
            image: PIL Image object
            
        Returns:
            Preprocessed PIL Image
        """
        try:
            # Convert to grayscale
            if image.mode != 'L':
                image = image.convert('L')
            
            # Resize if too small (OCR works better with larger images)
            width, height = image.size
            if max(width, height) < 300:
                scale_factor = 300 / max(width, height)
                new_size = (int(width * scale_factor), int(height * scale_factor))
                image = image.resize(new_size, Image.Resampling.LANCZOS)
            
            return image
            
        except Exception as e:
            logger.warning(f"Image preprocessing failed: {str(e)}")
            return image
    
    def is_supported_type(self, mime_type: str) -> bool:
        """
        Check if file type is supported for text extraction.
        
        Args:
            mime_type: MIME type to check
            
        Returns:
            True if supported, False otherwise
        """
        return mime_type.lower() in self.SUPPORTED_MIME_TYPES
    
    def get_file_type(self, mime_type: str) -> Optional[str]:
        """
        Get file type category from MIME type.
        
        Args:
            mime_type: MIME type
            
        Returns:
            File type string or None if unsupported
        """
        return self.SUPPORTED_MIME_TYPES.get(mime_type.lower())


class AttachmentProcessorWorker:
    """
    Async worker for processing attachments with queue management.
    """
    
    def __init__(self, processor: AttachmentProcessor, max_concurrent: int = 5):
        """
        Initialize attachment processor worker.
        
        Args:
            processor: AttachmentProcessor instance
            max_concurrent: Maximum concurrent processing tasks
        """
        self.processor = processor
        self.max_concurrent = max_concurrent
        self.semaphore = asyncio.Semaphore(max_concurrent)
        self.processing_count = 0
    
    async def process_attachment(
        self,
        file_content: bytes,
        filename: str,
        mime_type: str,
        attachment_id: str
    ) -> Dict[str, Any]:
        """
        Process a single attachment with concurrency control.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            mime_type: MIME type
            attachment_id: Unique attachment identifier
            
        Returns:
            Processing result with metadata
        """
        async with self.semaphore:
            self.processing_count += 1
            logger.info(f"Processing attachment {attachment_id} ({self.processing_count} active)")
            
            try:
                result = await self.processor.extract_text(file_content, filename, mime_type)
                result['attachment_id'] = attachment_id
                result['filename'] = filename
                result['mime_type'] = mime_type
                return result
                
            finally:
                self.processing_count -= 1
    
    async def process_batch(
        self,
        attachments: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        Process multiple attachments concurrently.
        
        Args:
            attachments: List of attachment dictionaries with keys:
                - file_content: bytes
                - filename: str
                - mime_type: str
                - attachment_id: str
                
        Returns:
            List of processing results
        """
        tasks = []
        for attachment in attachments:
            task = self.process_attachment(
                file_content=attachment['file_content'],
                filename=attachment['filename'],
                mime_type=attachment['mime_type'],
                attachment_id=attachment['attachment_id']
            )
            tasks.append(task)
        
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Handle exceptions in results
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                processed_results.append({
                    'attachment_id': attachments[i]['attachment_id'],
                    'filename': attachments[i]['filename'],
                    'status': 'failed',
                    'error': str(result),
                    'text': None
                })
            else:
                processed_results.append(result)
        
        return processed_results
